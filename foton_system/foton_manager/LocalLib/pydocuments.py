# Importações de bibliotecas necessárias
import json
import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

# Constantes
TABLE_STYLE = 'Table Grid'
FONT_SIZE = Pt(10)

# Funções

def insert_table(doc, table_data):
    """
    Insere uma tabela em um documento .docx.

    Args:
        doc (Document): O objeto Document do documento .docx.
        table_data (list): Uma lista de listas representando os dados da tabela.

    Returns:
        None
    """
    if not table_data:
        return

    try:
        table = doc.add_table(rows=1, cols=len(table_data[0]))
        table.autofit = False
        table.allow_autofit = False
        table.style = TABLE_STYLE

        for row_data in table_data:
            row = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                cell = row[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = FONT_SIZE

    except KeyError:
        # Se o estilo 'Table Grid' não estiver disponível, use um estilo padrão
        table.style = 'LightShading-Accent1'

def replace_text(doc, replacements):
    """
    Substitui o texto em um documento .docx de acordo com um dicionário de substituições.

    Args:
        doc (Document): O objeto Document do documento .docx.
        replacements (dict): Um dicionário de substituições no formato {texto_antigo: texto_novo}.

    Returns:
        None
    """
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(old_text, new_text)

def validate_args(input_path, output_path, file_name):
    """
    Valida os argumentos da função edit_and_save_docx.

    Args:
        input_path (str): O caminho do arquivo de entrada (.docx).
        output_path (str): O caminho do diretório de saída.
        file_name (str): O nome do arquivo de saída (sem extensão).

    Raises:
        FileNotFoundError: Se o arquivo de entrada não for encontrado.
    """
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Arquivo de entrada '{input_path}' não encontrado.")

def load_data_from_txt(txt_file_path):
    """
    Carrega dados de substituição de um arquivo de texto.

    Args:
        txt_file_path (str): O caminho do arquivo de texto.

    Returns:
        dict or None: Um dicionário de substituições ou None se o arquivo não for encontrado.
    """
    try:
        with open(txt_file_path, 'r') as txt_file:
            replacements = {}
            for line in txt_file:
                old_text, new_text = line.strip().split(';')
                replacements[old_text] = new_text
            return replacements
    except FileNotFoundError:
        return None

def load_data_from_json(json_file_path):
    """
    Carrega dados de uma tabela de um arquivo JSON.

    Args:
        json_file_path (str): O caminho do arquivo JSON.

    Returns:
        list or None: Uma lista de listas representando os dados da tabela ou None se o arquivo não for encontrado.
    """
    try:
        with open(json_file_path, 'r') as json_file:
            table_data = json.load(json_file)
            return table_data
    except FileNotFoundError:
        return None

def process_argument(arg):
    """
    Processa um argumento que pode ser um dicionário, uma lista ou um caminho para um arquivo.
    
    Args:
        arg: O argumento a ser processado.

    Returns:
        dict, list, or None: O argumento processado como um dicionário, lista ou None se não for um tipo suportado.
    """
    if isinstance(arg, dict) or isinstance(arg, list):
        return arg
    elif os.path.isfile(arg):
        # Se for um arquivo JSON, carrega os dados do arquivo
        if arg.lower().endswith('.json'):
            return load_data_from_json(arg)
        # Se for um arquivo de texto, carrega as substituições do arquivo
        elif arg.lower().endswith('.txt'):
            return load_data_from_txt(arg)
    return None

def edit_and_save_docx(input_path, output_path, file_name, replacements=None, table_data=None):
    """
    Edita um arquivo .docx com opção de substituir texto e inserir uma tabela,
    e salva a edição em um novo arquivo .docx. Também pode converter o documento para PDF.

    Args:
        input_path (str): O caminho do arquivo de entrada (.docx).
        output_path (str): O caminho do diretório de saída.
        file_name (str): O nome do arquivo de saída (sem extensão).
        replacements (dict, list, str, optional): Um dicionário de substituições, uma lista de listas representando os dados da tabela ou um caminho para um arquivo que contém os dados.
        table_data (dict, list, str, optional): Um dicionário de substituições, uma lista de listas representando os dados da tabela ou um caminho para um arquivo que contém os dados.

    Returns:
        None
    """
    try:
        # Validação de argumentos
        validate_args(input_path, output_path, file_name)

        # Abre o documento .docx
        doc = Document(input_path)

        # Processa o argumento replacements, se fornecido
        if replacements is not None:
            replacements = process_argument(replacements)

        # Processa o argumento table_data, se fornecido
        if table_data is not None:
            table_data = process_argument(table_data)

        # Substitui o texto, se houver dados de substituição
        if replacements:
            replace_text(doc, replacements)

        # Insere a tabela, se houver dados de tabela
        if table_data:
            insert_table(doc, table_data)

        # Salva o documento editado em um novo arquivo .docx
        output_docx = os.path.join(output_path, f"{file_name}.docx")
        doc.save(output_docx)
        print(f"Documento .docx editado salvo em: '{output_docx}'.")

        # Converte o documento para PDF, se necessário
        output_pdf = os.path.join(output_path, f"{file_name}.pdf")
        convert(output_docx, output_pdf)
        print(f"Documento PDF editado salvo em: '{output_pdf}'.")
    
    except FileNotFoundError as e:
        print(f"Erro: {str(e)} Arquivo não encontrado.")
    except Exception as e:
        print(f"Ocorreu um erro: {str(e)}")

# Exemplo de uso:

if __name__ == "__main__":
    path = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(path,'teste.docx')
    output_path = path
    file_name = "teste_editado"
    
    replacements = {
        "Texto_antigo1": "novo_texto1",
        "Texto_antigo2": "novo_texto2",
    }
    
    table_data = [
        ["Coluna1", "Coluna2", "Coluna3"],
        ["Dado1", "Dado2", "Dado3"],
    ]

    edit_and_save_docx(input_path, output_path, file_name, replacements, )
